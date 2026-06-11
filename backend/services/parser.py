import io
import logging
from pypdf import PdfReader
from docx import Document

logger = logging.getLogger("screener_parser")

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts all text from a PDF file."""
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        text_list = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_list.append(page_text)
            else:
                logger.warning(f"No text extracted from page {i} of PDF")
        
        extracted = "\n".join(text_list).strip()
        if not extracted:
            raise ValueError("Extracted text is empty. The PDF might contain only scanned images or be empty.")
        return extracted
    except Exception as e:
        logger.error(f"Failed to parse PDF: {e}")
        raise ValueError(f"Error parsing PDF: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts all text from a DOCX file, including paragraphs and tables."""
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        text_list = []
        
        # Read paragraphs
        for p in doc.paragraphs:
            if p.text:
                text_list.append(p.text)
                
        # Read tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    text_list.append(" | ".join(row_text))
                    
        extracted = "\n".join(text_list).strip()
        if not extracted:
            raise ValueError("Extracted text is empty. The DOCX file might be empty.")
        return extracted
    except Exception as e:
        logger.error(f"Failed to parse DOCX: {e}")
        raise ValueError(f"Error parsing DOCX: {str(e)}")

def sanitize_text(text: str) -> str:
    """Sanitizes raw text: strips formatting, removes duplicate spaces/newlines."""
    if not text:
        return ""
    # Strip unnecessary character patterns, convert multiple spacing to single
    lines = [line.strip() for line in text.splitlines()]
    # Remove empty lines
    non_empty_lines = [line for line in lines if line]
    return "\n".join(non_empty_lines)
